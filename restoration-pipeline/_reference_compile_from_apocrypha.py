#!/usr/bin/env python3
"""
Compile The Apocrypha Restored Names Edition to .docx
Uses python-docx. Run: python3 compile.py
"""
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.environ.get(
    'APOC_ROOT',
    os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
)
WORKING = os.path.join(ROOT, 'working')
SOURCE_DIR = os.path.join(WORKING, 'modernized')
COMMENTARY_DIR = os.path.join(WORKING, 'commentary')
OUTPUT = os.environ.get(
    'APOC_OUTPUT',
    os.path.join(ROOT, 'The-Apocrypha-Restored-Names-Edition.docx')
)

# Books in 1611 order
BOOKS = [
    ('01', '1 Esdras', '01-1_Esdras.txt', '01-1-esdras-commentary.md'),
    ('02', '2 Esdras', '02-2_Esdras.txt', '02-2-esdras-commentary.md'),
    ('03', 'Tobit', '03-Tobit.txt', '03-tobit-commentary.md'),
    ('04', 'Judith', '04-Judith.txt', '04-judith-commentary.md'),
    ('05', 'The Rest of Esther', '05-The_Rest_of_Esther.txt', '05-rest-of-esther-commentary.md'),
    ('06', 'The Wisdom of Solomon', '06-The_Wisdom_of_Solomon.txt', '06-wisdom-of-solomon-commentary.md'),
    ('07', 'Ecclesiasticus', '07-Ecclesiasticus.txt', '07-ecclesiasticus-commentary.md'),
    ('08', 'Baruch with the Letter of Jeremiah', '08-Baruch_with_the_Letter_of_Jeremiah.txt', '08-baruch-commentary.md'),
    ('09', 'The Song of the Three Holy Children', '09-The_Song_of_the_Three_Holy_Children.txt', '09-song-of-three-commentary.md'),
    ('10', 'The History of Susanna', '10-The_History_of_Susanna.txt', '10-susanna-commentary.md'),
    ('11', 'Bel and the Dragon', '11-Bel_and_the_Dragon.txt', '11-bel-and-dragon-commentary.md'),
    ('12', 'The Prayer of Manasseh', '12-The_Prayer_of_Manasseh.txt', '12-prayer-of-manasseh-commentary.md'),
    ('13', '1 Maccabees', '13-1_Maccabees.txt', '13-1-maccabees-commentary.md'),
    ('14', '2 Maccabees', '14-2_Maccabees.txt', '14-2-maccabees-commentary.md'),
]


def parse_source_text(text):
    """Parse source text into list of {num, verses}."""
    chapters = []
    current = None
    for raw in text.split('\n'):
        line = raw.strip()
        if not line:
            continue
        m = re.match(r'^\[CHAPTER (\d+)\]$', line)
        if m:
            current = {'num': int(m.group(1)), 'verses': []}
            chapters.append(current)
            continue
        if line.startswith('===') and line.endswith('==='):
            continue
        vm = re.match(r'^(\d+)\s+(.*)$', line)
        if vm and current is not None:
            current['verses'].append({'n': int(vm.group(1)), 'text': vm.group(2).strip()})
        elif current is not None and current['verses']:
            current['verses'][-1]['text'] += ' ' + line
    return chapters


def parse_commentary(text):
    """Parse commentary md into {chapter_num: text} plus __preface."""
    sections = {}
    lines = text.split('\n')
    current_num = None
    buffer = []
    preface = []
    in_preface = True
    for line in lines:
        m = re.match(r'^## Chapter (\d+)', line, re.IGNORECASE)
        if m:
            if current_num is not None:
                sections[current_num] = '\n'.join(buffer).strip()
            elif in_preface:
                sections['__preface'] = '\n'.join(preface).strip()
                in_preface = False
            current_num = int(m.group(1))
            buffer = []
            continue
        if current_num is None and in_preface:
            preface.append(line)
        else:
            buffer.append(line)
    if current_num is not None:
        sections[current_num] = '\n'.join(buffer).strip()
    elif in_preface:
        sections['__preface'] = '\n'.join(preface).strip()
    return sections


def parse_book_intros(text):
    """Parse intros file into {title: body}."""
    intros = {}
    parts = re.split(r'^## Introduction to ', text, flags=re.MULTILINE)[1:]
    for part in parts:
        nl = part.find('\n')
        title = part[:nl].strip()
        body = part[nl+1:].strip()
        intros[title] = body
    return intros


def add_paragraph_with_inline_italics(doc, text, style=None, italic_block=False, indent_left=None):
    """Add paragraph with parsing for **bold** and *italic* runs.
    Bold is matched before italic so that ``**foo**`` does not collapse into
    a stray-asterisk + italicised middle.
    """
    p = doc.add_paragraph()
    if style:
        p.style = doc.styles[style]
    if indent_left is not None:
        p.paragraph_format.left_indent = Inches(indent_left)
        p.paragraph_format.right_indent = Inches(indent_left)
    if italic_block:
        run = p.add_run(text)
        run.italic = True
        return p
    # Match **bold** (greedy non-asterisk) OR *italic*.
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*)')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            p.add_run(text[last:m.start()])
        chunk = m.group(0)
        if chunk.startswith('**') and chunk.endswith('**') and len(chunk) > 4:
            run = p.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith('*') and chunk.endswith('*') and len(chunk) > 2:
            run = p.add_run(chunk[1:-1])
            run.italic = True
        else:
            p.add_run(chunk)
        last = m.end()
    if last < len(text):
        p.add_run(text[last:])
    return p


def render_markdown(doc, md, h1_page_break=True, skip_title=False):
    """Render markdown to docx paragraphs."""
    lines = md.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith('# '):
            if not skip_title:
                if h1_page_break:
                    p = doc.add_paragraph()
                    p.add_run().add_break(WD_BREAK.PAGE)
                doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith('---'):
            i += 1
            continue
        # Collect paragraph
        buf = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith('#') and not lines[i].startswith('---'):
            buf.append(lines[i])
            i += 1
        text = ' '.join(b.strip() for b in buf)
        text = re.sub(r'\s+', ' ', text).strip()
        # Italic block (paragraph wrapped in *...*)
        if re.match(r'^\*[^*]+\*$', text):
            add_paragraph_with_inline_italics(
                doc, text[1:-1], italic_block=True, indent_left=0.4
            )
        else:
            add_paragraph_with_inline_italics(doc, text)


def add_chapter_text(doc, chapter):
    """Add scripture text for a chapter as italicized indented paragraphs."""
    # Group verses ~5 at a time
    group = []
    for v in chapter['verses']:
        group.append(f"{v['n']} {v['text']}")
        if len(group) >= 5:
            text = ' '.join(group)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.25)
            run = p.add_run(text)
            run.italic = True
            group = []
    if group:
        text = ' '.join(group)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.25)
        run = p.add_run(text)
        run.italic = True


def add_divider(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('⸻')
    run.font.size = Pt(14)


def setup_document():
    doc = Document()

    # Page size: 7×10 KDP
    section = doc.sections[0]
    section.page_width = Inches(7)
    section.page_height = Inches(10)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # Default style: Cambria 11pt
    style = doc.styles['Normal']
    style.font.name = 'Cambria'
    style.font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_after = Pt(8)
    pf.line_spacing = 1.4

    # Customize headings
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Cambria'
    h1.font.size = Pt(20)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(12)

    h2 = doc.styles['Heading 2']
    h2.font.name = 'Cambria'
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(9)

    h3 = doc.styles['Heading 3']
    h3.font.name = 'Cambria'
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.italic = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.space_before = Pt(9)
    h3.paragraph_format.space_after = Pt(6)

    # Add page numbers in footer
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run('— ')
    run.font.size = Pt(9)
    # Add PAGE field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    pageRun = fp.add_run()
    pageRun._r.append(fldChar1)
    pageRun._r.append(instrText)
    pageRun._r.append(fldChar2)
    pageRun.font.size = Pt(9)
    runEnd = fp.add_run(' —')
    runEnd.font.size = Pt(9)

    return doc


def main():
    doc = setup_document()

    # --- Title page ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('\n\n\n')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('The Apocrypha')
    run.font.size = Pt(40)
    run.font.bold = True
    run.font.name = 'Cambria'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Restored Names Edition')
    run.font.size = Pt(22)
    run.font.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('\n')
    run = p.add_run('With Complete Nazarite Commentary, Cross-References, and Restored Divine Names')
    run.font.size = Pt(13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('\n\n\n\n\n')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('By Yoshi')
    run.font.size = Pt(16)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Remnant of Promise')
    run.font.size = Pt(13)

    # Page break after title
    doc.add_page_break()

    # --- Copyright page ---
    p = doc.add_paragraph()
    p.add_run('The Apocrypha — Restored Names Edition').font.bold = True
    doc.add_paragraph('With Complete Nazarite Commentary, Cross-References, and Restored Divine Names')
    doc.add_paragraph('By Yoshi')
    doc.add_paragraph('Published by Remnant of Promise — remnantofpromise.org')
    doc.add_paragraph('')
    doc.add_paragraph(
        'The fourteen books of the King James Apocrypha as set in the original 1611 King '
        'James Bible. The source text is the public-domain 1769 Cambridge / Blayney '
        'standardized revision of the King James Apocrypha, drawn from the public-domain '
        'digital release of the King James Bible with Apocrypha distributed by eBible.org '
        'under public-domain status.'
    )
    doc.add_paragraph('')
    doc.add_paragraph(
        'Editorial restoration of the divine name (YAHUAH where the textual evidence supports '
        'the Tetragrammaton), modernization of archaic English to contemporary forms, light '
        'internal punctuation, and the Nazarite commentary herein are original to this '
        'edition. The chapter and verse divisions of the 1769 King James text are preserved '
        'exactly. No verses have been added or removed. No passage has been paraphrased.'
    )
    doc.add_paragraph('')
    doc.add_paragraph(
        'All scripture witness is offered with the prayer that the Spirit will use it for the '
        'awakening of the remnant of YAHUAH.'
    )

    doc.add_page_break()

    # --- Front matter ---
    with open(os.path.join(WORKING, '00-front-matter.md'), 'r') as f:
        fm = f.read()
    # Skip the title block; start at first "# Introduction"
    idx = fm.find('# Introduction')
    if idx >= 0:
        fm = fm[idx:]
    render_markdown(doc, fm)

    # --- Book intros ---
    with open(os.path.join(WORKING, '01-book-intros.md'), 'r') as f:
        intros_text = f.read()
    intros = parse_book_intros(intros_text)
    intro_key_map = {
        '1 Esdras': '1 Esdras',
        '2 Esdras': '2 Esdras',
        'Tobit': 'Tobit',
        'Judith': 'Judith',
        'The Rest of Esther': 'The Rest of Esther',
        'The Wisdom of Solomon': 'The Wisdom of Solomon',
        'Ecclesiasticus (The Wisdom of Jesus the Son of Sirach)': 'Ecclesiasticus',
        'Baruch with the Letter of Jeremiah': 'Baruch with the Letter of Jeremiah',
        'The Song of the Three Holy Children': 'The Song of the Three Holy Children',
        'The History of Susanna': 'The History of Susanna',
        'Bel and the Dragon': 'Bel and the Dragon',
        'The Prayer of Manasseh': 'The Prayer of Manasseh',
        '1 Maccabees': '1 Maccabees',
        '2 Maccabees': '2 Maccabees',
    }

    # --- Iterate books ---
    # Each scripture chapter and each book's introductory page is given its
    # own Heading 1 with a disambiguating title. Vellum starts a new chapter
    # at every Heading 1, and (per the yoshi-voice skill rule) thematic
    # grouping is achieved by repeating the section name across consecutive
    # chapter titles, not by inserting standalone section dividers.
    for num, title, src_file, com_file in BOOKS:
        # ---- Book intro chapter (Heading 1: "{title} — Introduction") ----
        doc.add_page_break()
        intro_heading = doc.add_heading(f'{title} — Introduction', level=1)
        intro_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Per-book intro body
        intro_key = next((k for k, v in intro_key_map.items() if v == title), None)
        if intro_key and intro_key in intros:
            render_markdown(doc, intros[intro_key], h1_page_break=False)

        # Source text + commentary
        with open(os.path.join(SOURCE_DIR, src_file), 'r') as f:
            src_text = f.read()
        chapters = parse_source_text(src_text)

        commentary = {}
        com_path = os.path.join(COMMENTARY_DIR, com_file)
        if os.path.exists(com_path):
            with open(com_path, 'r') as f:
                commentary = parse_commentary(f.read())

        # Render preface (if commentary has a __preface block) into the
        # intro chapter, just below the per-book intro body.
        if '__preface' in commentary and commentary['__preface']:
            # Strip the preface's own H1 (it is the book title, already
            # carried by the chapter heading above).
            pre = commentary['__preface']
            pre_lines = pre.split('\n')
            cleaned = []
            for line in pre_lines:
                if line.startswith('# '):
                    continue
                cleaned.append(line)
            cleaned_pre = '\n'.join(cleaned).strip()
            if cleaned_pre:
                render_markdown(doc, cleaned_pre, h1_page_break=False)

        # ---- Each scripture chapter as its own Heading 1 ----
        for chapter in chapters:
            doc.add_page_break()
            ch_heading = doc.add_heading(
                f'{title} — Chapter {chapter["num"]}', level=1
            )
            ch_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_chapter_text(doc, chapter)
            add_divider(doc)
            com = commentary.get(chapter['num'])
            if com:
                # Strip leading parenthetical heading lines
                com_lines = com.split('\n')
                while com_lines and (not com_lines[0].strip() or re.match(r'^\(.*\)$', com_lines[0].strip())):
                    com_lines.pop(0)
                cleaned = '\n'.join(com_lines)
                doc.add_heading('Commentary', level=2)
                render_markdown(doc, cleaned, h1_page_break=False)

    # --- Closing essay ---
    doc.add_page_break()
    with open(os.path.join(WORKING, '02-closing-essay.md'), 'r') as f:
        render_markdown(doc, f.read(), h1_page_break=False)

    # --- Appendices ---
    doc.add_page_break()
    with open(os.path.join(WORKING, '03-appendices.md'), 'r') as f:
        render_markdown(doc, f.read(), h1_page_break=False)

    # --- About the Author ---
    doc.add_page_break()
    with open(os.path.join(WORKING, '04-about-the-author.md'), 'r') as f:
        render_markdown(doc, f.read(), h1_page_break=False)

    # Save
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    size_mb = os.path.getsize(OUTPUT) / 1024 / 1024
    print(f"Wrote: {OUTPUT}")
    print(f"Size: {size_mb:.2f} MB")


if __name__ == '__main__':
    main()
